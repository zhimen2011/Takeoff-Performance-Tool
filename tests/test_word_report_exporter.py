from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"
sys.path.insert(0, str(SRC_DIR))

from stas_app.exporters.word_report import WordReportExporter, split_stas_report_sections


SAMPLE_OUTPUT = ROOT_DIR / "tests" / "fixtures" / "STASOUT_SAMPLE.out"
TEMPLATE_DIR = ROOT_DIR / "templates" / "reports" / "manual_takeoff"


class WordReportExporterTests(unittest.TestCase):
    def test_split_sections_starts_at_first_elevation(self) -> None:
        content = SAMPLE_OUTPUT.read_text(encoding="utf-8")

        sections = split_stas_report_sections(content)

        self.assertEqual(len(sections), 2)
        self.assertTrue(sections[0].lstrip().startswith("ELEVATION 1000 FT"))
        self.assertNotIn("HEADER TEXT", sections[0])

    def test_split_sections_removes_no_emergency_turn_procedure_body(self) -> None:
        content = SAMPLE_OUTPUT.read_text(encoding="utf-8")

        sections = split_stas_report_sections(content)

        self.assertNotIn("THIS LINE SHOULD NOT BE INCLUDED", sections[0])
        self.assertIn("TURN RIGHT CLIMBING", sections[1])

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_generates_docx(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "STASOUT.docx"

            result = WordReportExporter().export(SAMPLE_OUTPUT, target)

            self.assertTrue(result.succeeded, result.error_message)
            self.assertTrue(target.exists())

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_uses_temporary_report_template(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "STASOUT.docx"

            result = WordReportExporter(TEMPLATE_DIR).export(SAMPLE_OUTPUT, target)

            self.assertTrue(result.succeeded, result.error_message)
            doc = Document(str(target))
            header_text = "\n".join(paragraph.text for paragraph in doc.sections[0].header.paragraphs)
            self.assertIn("临时起飞分析", header_text)
            self.assertIn("STAS Temporary Body", [style.name for style in doc.styles])

    def test_missing_source_returns_error(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "missing.docx"

            result = WordReportExporter().export(Path(temp_dir) / "missing.out", target)

            self.assertFalse(result.succeeded)
            self.assertIn("does not exist", result.error_message)

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_replaces_report_date_in_docx_body(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "STASOUT.out"
            target = Path(temp_dir) / "STASOUT.docx"
            source.write_text(
                "ELEVATION 1000 FT\n      738         CFM56-7B26                                DATED 19-MAY-2026\n",
                encoding="utf-8",
            )

            result = WordReportExporter(TEMPLATE_DIR).export(source, target, report_date_override="04-APR-2026")

            self.assertTrue(result.succeeded, result.error_message)
            paragraphs = [paragraph.text for paragraph in Document(str(target)).paragraphs]
            self.assertTrue(any("DATED 04-APR-2026" in paragraph for paragraph in paragraphs))
            self.assertFalse(any("DATED 19-MAY-2026" in paragraph for paragraph in paragraphs))
            self.assertIn("DATED 19-MAY-2026", source.read_text(encoding="utf-8"))


if __name__ == "__main__":
    unittest.main()
