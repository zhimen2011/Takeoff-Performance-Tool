from __future__ import annotations

import importlib.util
import sys
import tempfile
import unittest
from pathlib import Path

ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"
sys.path.insert(0, str(SRC_DIR))

from stas_app.exporters.manual_takeoff_report import (
    MANUAL_SECTION_COMPACT_LAYOUT,
    MANUAL_SECTION_NORMAL_LAYOUT,
    MANUAL_SECTION_TIGHT_LAYOUT,
    ManualTakeoffReportExporter,
    manual_section_layout,
    split_manual_takeoff_report_sections,
)
from stas_app.models.request import PerformanceRequest


SAMPLE_OUTPUT = ROOT_DIR / "tests" / "fixtures" / "STASOUT_SAMPLE.out"
TEMPLATE_DIR = ROOT_DIR / "templates" / "reports" / "manual_takeoff"


class ManualTakeoffReportExporterTests(unittest.TestCase):
    def test_templates_are_loaded_from_registry_config(self) -> None:
        exporter = ManualTakeoffReportExporter(TEMPLATE_DIR)

        templates = exporter.templates()

        self.assertEqual([template.id for template in templates], ["738_normal", "777f_normal", "777f_derate", "777f_bump"])

    def test_split_manual_sections_preserves_manual_report_indent(self) -> None:
        content = SAMPLE_OUTPUT.read_text(encoding="utf-8")

        sections = split_manual_takeoff_report_sections(content)

        self.assertTrue(sections)
        self.assertTrue(sections[0].startswith("      ELEVATION 1000 FT"))

    def test_manual_section_layout_uses_tighter_spacing_for_long_sections(self) -> None:
        normal_section = "\n".join(f"LINE {index}" for index in range(10))
        compact_section = "\n".join(f"LINE {index}" for index in range(64))
        tight_section = "\n".join(f"LINE {index}" for index in range(72))

        self.assertEqual(manual_section_layout(normal_section), MANUAL_SECTION_NORMAL_LAYOUT)
        self.assertEqual(manual_section_layout(compact_section), MANUAL_SECTION_COMPACT_LAYOUT)
        self.assertEqual(manual_section_layout(tight_section), MANUAL_SECTION_TIGHT_LAYOUT)

    def test_incompatible_template_returns_error(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "manual.docx"

            result = ManualTakeoffReportExporter(TEMPLATE_DIR).export(
                SAMPLE_OUTPUT,
                target,
                "777f_bump",
                PerformanceRequest("738", "ZBAA"),
            )

            self.assertFalse(result.succeeded)
            self.assertIn("不适用于机型", result.error_message)
            self.assertFalse(target.exists())

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_generates_docx_from_template(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "manual.docx"

            result = ManualTakeoffReportExporter(TEMPLATE_DIR).export(
                SAMPLE_OUTPUT,
                target,
                "738_normal",
                PerformanceRequest("738", "ZBAA"),
            )

            self.assertTrue(result.succeeded, result.error_message)
            self.assertTrue(target.exists())

            from docx import Document

            doc = Document(str(target))
            self.assertTrue(doc.paragraphs[0].text.startswith("      ELEVATION 1000 FT"))
            self.assertTrue(doc.paragraphs[0].paragraph_format.keep_with_next)
            self.assertTrue(doc.paragraphs[2].paragraph_format.page_break_before)

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_replaces_report_date_without_modifying_source(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "STASOUT.out"
            target = Path(temp_dir) / "manual.docx"
            source.write_text(
                "ELEVATION 1000 FT\n      738         CFM56-7B26                                DATED 19-MAY-2026\n",
                encoding="utf-8",
            )

            result = ManualTakeoffReportExporter(TEMPLATE_DIR).export(
                source,
                target,
                "738_normal",
                PerformanceRequest("738", "ZBAA", report_date_override="04-APR-2026"),
            )

            self.assertTrue(result.succeeded, result.error_message)

            from docx import Document

            texts = [paragraph.text for paragraph in Document(str(target)).paragraphs]
            self.assertTrue(any("DATED 04-APR-2026" in text for text in texts))
            self.assertIn("DATED 19-MAY-2026", source.read_text(encoding="utf-8"))

    @unittest.skipIf(importlib.util.find_spec("docx") is None, "python-docx is not installed")
    def test_export_highlights_manual_keywords_without_changing_text(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            target = Path(temp_dir) / "manual.docx"
            sections = [
                "\n".join(
                    (
                        "      ELEVATION 1000 FT",
                        "      777F         GE90-110B1L                  10% DERATE   DATED 16-APR-2026",
                        "      *** FLAPS 15 ***   AIR COND ON     ANTI-ICE ENG ONLY        SHUANGLIU",
                        "      777F         GE90-110B1L                  20% DERATE   DATED 16-APR-2026",
                    )
                )
            ]

            result = ManualTakeoffReportExporter(TEMPLATE_DIR).export_sections(
                sections,
                target,
                "777f_derate",
                (PerformanceRequest("777F", "ZUUU", thrust_option="减推力10%"),),
            )

            self.assertTrue(result.succeeded, result.error_message)

            from docx import Document
            from docx.enum.text import WD_COLOR_INDEX

            doc = Document(str(target))
            texts = [paragraph.text for paragraph in doc.paragraphs]
            self.assertIn(sections[0].splitlines()[1], texts)
            self.assertIn(sections[0].splitlines()[2], texts)
            self.assertIn(sections[0].splitlines()[3], texts)
            self.assertEqual(_highlighted_text(doc, WD_COLOR_INDEX.YELLOW), ["10% DERATE", "ANTI-ICE ENG ONLY"])


def _highlighted_text(doc, highlight_color) -> list[str]:
    highlighted: list[str] = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:
                highlighted.append(run.text)
    return highlighted


if __name__ == "__main__":
    unittest.main()
