"""Manual takeoff analysis Word report generation."""

from __future__ import annotations

import time
from collections.abc import Sequence
from dataclasses import dataclass
from pathlib import Path

from stas_app.exporters.word_report import split_stas_report_sections
from stas_app.models.manual_report import ManualTakeoffReportTemplate, ManualTakeoffTemplateRegistry
from stas_app.models.report import ReportExportResult
from stas_app.models.request import PerformanceRequest
from stas_app.utils.aviation_date import replace_report_date


@dataclass(frozen=True)
class ManualSectionLayout:
    """Font and line spacing used for one manual report page."""

    font_size_pt: float
    line_spacing_pt: float


MANUAL_SECTION_NORMAL_LAYOUT = ManualSectionLayout(font_size_pt=7.5, line_spacing_pt=9.3)
MANUAL_SECTION_COMPACT_LAYOUT = ManualSectionLayout(font_size_pt=7.2, line_spacing_pt=8.4)
MANUAL_SECTION_TIGHT_LAYOUT = ManualSectionLayout(font_size_pt=6.8, line_spacing_pt=7.4)
MANUAL_SECTION_CHARACTERS_PER_LINE = 108
MANUAL_SECTION_COMPACT_LINE_THRESHOLD = 62
MANUAL_SECTION_TIGHT_LINE_THRESHOLD = 70
MANUAL_HIGHLIGHT_TERMS = ("ANTI-ICE ENG ONLY", "10% DERATE")


class ManualTakeoffReportExporter:
    """Generate the manual takeoff analysis format from STAS raw output."""

    def __init__(
        self,
        template_dir: str | Path = "templates/reports/manual_takeoff",
        registry: ManualTakeoffTemplateRegistry | None = None,
    ) -> None:
        self.template_dir = Path(template_dir)
        self.registry = registry or ManualTakeoffTemplateRegistry.from_directory(self.template_dir)

    def templates(self) -> tuple[ManualTakeoffReportTemplate, ...]:
        return self.registry.all()

    def export(
        self,
        stas_output_path: str | Path,
        docx_path: str | Path,
        template_id: str,
        request: PerformanceRequest | None = None,
    ) -> ReportExportResult:
        """Generate a manual-format Word report from one raw STAS output file."""

        started = time.perf_counter()
        source_path = Path(stas_output_path)
        target_path = Path(docx_path)

        if not source_path.exists():
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"STAS output file does not exist: {source_path}",
            )

        try:
            sections = split_manual_takeoff_report_sections(source_path.read_text(encoding="utf-8", errors="replace"))
            return self.export_sections(sections, target_path, template_id, (request,) if request else ())
        except Exception as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"Manual takeoff Word report generation failed: {exc}",
            )

    def export_sections(
        self,
        sections: Sequence[str],
        docx_path: str | Path,
        template_id: str,
        requests: Sequence[PerformanceRequest] = (),
    ) -> ReportExportResult:
        """Generate a manual-format Word report from pre-split report sections."""

        started = time.perf_counter()
        target_path = Path(docx_path)
        report_sections = tuple(section.rstrip("\n") for section in sections if section.strip())
        if not report_sections:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message="No manual takeoff report sections were provided",
            )

        try:
            template = self.registry.get(template_id)
            validation_error = self._validate_requests(template.id, requests)
            if validation_error:
                return ReportExportResult(
                    status="error",
                    output_path=target_path,
                    elapsed_seconds=time.perf_counter() - started,
                    error_message=validation_error,
                )

            template_path = self.template_dir / template.filename
            if not template_path.exists():
                return ReportExportResult(
                    status="error",
                    output_path=target_path,
                    elapsed_seconds=time.perf_counter() - started,
                    error_message=f"Manual takeoff report template does not exist: {template_path}",
                )

            report_sections = _apply_report_date_override(report_sections, requests)
            self._generate(report_sections, template_path, target_path)
        except ImportError as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"python-docx is required to generate manual takeoff Word reports: {exc}",
            )
        except Exception as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"Manual takeoff Word report generation failed: {exc}",
            )

        return ReportExportResult(
            status="success",
            output_path=target_path,
            elapsed_seconds=time.perf_counter() - started,
        )

    def _validate_requests(self, template_id: str, requests: Sequence[PerformanceRequest]) -> str:
        for request in requests:
            error = self.registry.validate_request(template_id, request)
            if error:
                return error
        return ""

    def _generate(self, sections: Sequence[str], template_path: Path, target_path: Path) -> None:
        from docx import Document
        from docx.enum.text import WD_COLOR_INDEX
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Pt

        target_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document(str(template_path))
        self._clear_body(doc, qn)

        for index, section_text in enumerate(sections):
            lines = section_text.splitlines()
            layout = manual_section_layout(section_text)
            for line_index, line in enumerate(lines):
                paragraph = doc.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                paragraph_format.line_spacing = Pt(layout.line_spacing_pt)
                paragraph_format.space_before = Pt(0)
                paragraph_format.space_after = Pt(0)
                paragraph_format.keep_together = True
                paragraph_format.keep_with_next = line_index < len(lines) - 1
                if index > 0 and line_index == 0:
                    paragraph_format.page_break_before = True

                self._add_highlighted_line(paragraph, line, layout, Pt, qn, WD_COLOR_INDEX)

        doc.save(str(target_path))

    def _add_highlighted_line(self, paragraph, line: str, layout: ManualSectionLayout, Pt, qn, color_index) -> None:
        cursor = 0
        for start, end in _manual_highlight_ranges(line):
            if start > cursor:
                self._add_manual_run(paragraph, line[cursor:start], layout, Pt, qn)
            self._add_manual_run(paragraph, line[start:end], layout, Pt, qn, color_index.YELLOW)
            cursor = end

        if cursor < len(line) or not line:
            self._add_manual_run(paragraph, line[cursor:], layout, Pt, qn)

    def _add_manual_run(self, paragraph, text: str, layout: ManualSectionLayout, Pt, qn, highlight=None) -> None:
        run = paragraph.add_run(text)
        run.font.name = "DengXian Light"
        run.font.size = Pt(layout.font_size_pt)
        if highlight is not None:
            run.font.highlight_color = highlight
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线 Light")

    def _clear_body(self, doc, qn) -> None:
        body = doc._body._element
        section_properties_tag = qn("w:sectPr")
        for child in list(body):
            if child.tag == section_properties_tag:
                continue
            body.remove(child)


def split_manual_takeoff_report_sections(content: str) -> list[str]:
    """Split raw STAS output into manual-report sections."""

    sections = [section.rstrip("\n") for section in split_stas_report_sections(content) if section.strip()]
    if sections:
        return sections
    return [content.rstrip("\n")] if content.strip() else []


def manual_section_layout(section_text: str) -> ManualSectionLayout:
    """Choose a compact manual-report layout so one section fits one page more often."""

    visual_line_count = _estimated_visual_line_count(section_text)
    if visual_line_count >= MANUAL_SECTION_TIGHT_LINE_THRESHOLD:
        return MANUAL_SECTION_TIGHT_LAYOUT
    if visual_line_count >= MANUAL_SECTION_COMPACT_LINE_THRESHOLD:
        return MANUAL_SECTION_COMPACT_LAYOUT
    return MANUAL_SECTION_NORMAL_LAYOUT


def _estimated_visual_line_count(section_text: str) -> int:
    visual_line_count = 0
    for line in section_text.splitlines() or [""]:
        length = len(line.rstrip("\n"))
        wrapped_lines = max(1, (length + MANUAL_SECTION_CHARACTERS_PER_LINE - 1) // MANUAL_SECTION_CHARACTERS_PER_LINE)
        visual_line_count += wrapped_lines
    return visual_line_count


def _manual_highlight_ranges(line: str) -> list[tuple[int, int]]:
    ranges: list[tuple[int, int]] = []
    for term in MANUAL_HIGHLIGHT_TERMS:
        start = 0
        while True:
            index = line.find(term, start)
            if index == -1:
                break
            ranges.append((index, index + len(term)))
            start = index + len(term)

    ranges.sort()
    merged: list[tuple[int, int]] = []
    for start, end in ranges:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return merged


def _apply_report_date_override(
    sections: Sequence[str],
    requests: Sequence[PerformanceRequest],
) -> tuple[str, ...]:
    report_date = _shared_report_date_override(requests)
    if not report_date:
        return tuple(sections)
    return tuple(replace_report_date(section, report_date) for section in sections)


def _shared_report_date_override(requests: Sequence[PerformanceRequest]) -> str:
    for request in requests:
        token = request.report_date_override.strip()
        if token:
            return token
    return ""
