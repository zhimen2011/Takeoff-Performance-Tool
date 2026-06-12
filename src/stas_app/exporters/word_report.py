"""Word report generation for STAS raw output files."""

from __future__ import annotations

import time
from collections.abc import Sequence
from pathlib import Path

from stas_app.models.report import ReportExportResult
from stas_app.models.temporary_report import TemporaryTakeoffTemplateRegistry
from stas_app.utils.aviation_date import replace_report_date


TEMPORARY_REPORT_TEMPLATE_ID = "temporary_default"
TEMPORARY_REPORT_BODY_STYLE = "STAS Temporary Body"


class WordReportExporter:
    """Generate a formatted Word report from STAS raw text output."""

    def __init__(
        self,
        template_dir: str | Path = "templates/reports/manual_takeoff",
        registry: TemporaryTakeoffTemplateRegistry | None = None,
        template_id: str = TEMPORARY_REPORT_TEMPLATE_ID,
    ) -> None:
        self.template_dir = Path(template_dir)
        self.registry = registry
        self.template_id = template_id

    def export(
        self,
        stas_output_path: str | Path,
        docx_path: str | Path,
        logo_path: str | Path | None = None,
        report_date_override: str = "",
    ) -> ReportExportResult:
        started = time.perf_counter()
        source_path = Path(stas_output_path)
        target_path = Path(docx_path)

        if not source_path.exists():
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"STAS output file does not exist: {source_path}",
            )

        try:
            self._generate(source_path, target_path, Path(logo_path) if logo_path else None, report_date_override)
        except ImportError as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"python-docx is required to generate Word reports: {exc}",
            )
        except Exception as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"Word report generation failed: {exc}",
            )

        return ReportExportResult(
            status="success",
            output_path=target_path,
            elapsed_seconds=time.perf_counter() - started,
        )

    def export_sections(
        self,
        sections: Sequence[str],
        docx_path: str | Path,
        logo_path: str | Path | None = None,
        report_date_override: str = "",
    ) -> ReportExportResult:
        """Generate a formatted Word report from pre-split report sections."""

        started = time.perf_counter()
        target_path = Path(docx_path)
        report_sections = tuple(
            replace_report_date(section.rstrip("\n"), report_date_override)
            for section in sections
            if section.strip()
        )
        if not report_sections:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message="No report sections were provided",
            )

        try:
            self._generate_sections(report_sections, target_path, Path(logo_path) if logo_path else None)
        except ImportError as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"python-docx is required to generate Word reports: {exc}",
            )
        except Exception as exc:
            return ReportExportResult(
                status="error",
                output_path=target_path,
                elapsed_seconds=time.perf_counter() - started,
                error_message=f"Word report generation failed: {exc}",
            )

        return ReportExportResult(
            status="success",
            output_path=target_path,
            elapsed_seconds=time.perf_counter() - started,
        )

    def _generate(
        self,
        source_path: Path,
        target_path: Path,
        logo_path: Path | None,
        report_date_override: str = "",
    ) -> None:
        content = source_path.read_text(encoding="utf-8", errors="replace")
        content = replace_report_date(content, report_date_override)
        self._generate_sections(
            split_stas_report_sections(content),
            target_path,
            logo_path,
        )

    def _generate_sections(self, sections: Sequence[str], target_path: Path, logo_path: Path | None) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt

        target_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document(str(self._template_path()))

        self._clear_body(doc, qn)
        self._write_body(
            doc=doc,
            sections=list(sections),
            inches=Inches,
            point=Pt,
            paragraph_alignment=WD_ALIGN_PARAGRAPH,
            line_spacing_rule=WD_LINE_SPACING,
        )

        doc.save(str(target_path))

    def _write_body(
        self,
        doc,
        sections: list[str],
        inches,
        point,
        paragraph_alignment,
        line_spacing_rule,
    ) -> None:
        body_style = _optional_style(doc, TEMPORARY_REPORT_BODY_STYLE)
        for index, section_text in enumerate(sections):
            if index > 0:
                doc.add_page_break()

            for line in section_text.splitlines():
                paragraph = doc.add_paragraph(style=body_style) if body_style else doc.add_paragraph()
                run = paragraph.add_run(line)
                if body_style is None:
                    run.font.name = "Consolas"
                    run.font.size = point(8.5)
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing_rule = line_spacing_rule.EXACTLY
                    paragraph_format.line_spacing = point(11)
                    paragraph_format.space_before = point(0)
                    paragraph_format.space_after = point(0)
                    paragraph_format.left_indent = inches(0.5)
                    paragraph.alignment = paragraph_alignment.LEFT

    def _template_path(self) -> Path:
        template = self._template_registry().get(self.template_id)
        template_path = self.template_dir / template.filename
        if not template_path.exists():
            raise FileNotFoundError(f"Temporary takeoff report template does not exist: {template_path}")
        return template_path

    def _template_registry(self) -> TemporaryTakeoffTemplateRegistry:
        if self.registry is None:
            self.registry = TemporaryTakeoffTemplateRegistry.from_directory(self.template_dir)
        return self.registry

    def _clear_body(self, doc, qn) -> None:
        body = doc._body._element
        section_properties_tag = qn("w:sectPr")
        for child in list(body):
            if child.tag == section_properties_tag:
                continue
            body.remove(child)


def split_stas_report_sections(content: str) -> list[str]:
    """Split raw STAS output into report sections using ELEVATION markers."""

    first_elevation = content.find("ELEVATION")
    if first_elevation == -1:
        return [content] if content.strip() else []

    relevant_content = content[first_elevation:]
    parts = relevant_content.split("ELEVATION")
    sections: list[str] = []
    for part in parts[1:]:
        section_text = _normalize_elevation_section(part)
        if section_text.strip():
            sections.append(section_text)

    return sections


def _normalize_elevation_section(part: str) -> str:
    if _should_skip_engine_out_section(part):
        part = part.split("ENG-OUT PROCEDURE:", 1)[0].rstrip("\n")

    return "      ELEVATION" + part.rstrip("\n")


def _should_skip_engine_out_section(part: str) -> bool:
    marker = "ENG-OUT PROCEDURE:"
    marker_index = part.find(marker)
    if marker_index == -1:
        return False

    remaining_text = part[marker_index + len(marker) :]
    lines = remaining_text.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        return "NO EMERGENCY TURN" in stripped

    return False


def _optional_style(doc, style_name: str):
    try:
        return doc.styles[style_name]
    except KeyError:
        return None
